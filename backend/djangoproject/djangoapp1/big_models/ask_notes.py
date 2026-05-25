import json
import os
import re
import requests
from pypdf import PdfReader
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# -------------------------- 核心配置 --------------------------
OLLAMA_BASE_URL = "http://127.0.0.1:11434"
EMBEDDING_MODEL = "nomic-embed-text:latest"  # 向量模型
CHAT_MODEL = "qwen2.5:7b"  # 对话模型（可换成 llama3, gemma 等）


# ----------------------------------------------------------------

# ====================== 1. 文件读取函数（和之前一致） ======================
def read_text_file(file_path: str) -> str:
    """读取 txt / md 文件"""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except:
        return ""


def read_pdf_file(file_path: str) -> str:
    """读取 PDF 文件"""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    except:
        return ""


def read_docx_file(file_path: str) -> str:
    """读取 Word 文件（含段落和表格）"""
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)
    except Exception as e:
        print(f"docx 读取失败 ({file_path}): {e}")
        return ""


def get_file_content(file_path: str) -> str:
    """自动识别文件类型"""
    if not os.path.exists(file_path):
        return ""
    ext = os.path.splitext(file_path)[-1].lower()
    if ext in [".txt", ".md"]:
        return read_text_file(file_path)
    elif ext == ".pdf":
        return read_pdf_file(file_path)
    elif ext in [".docx", ".doc"]:
        return read_docx_file(file_path)
    return ""


# ====================== 2. Ollama Embedding 向量生成 ======================
def get_embedding(text: str) -> list:
    """调用 Ollama 生成向量"""
    if not text:
        return []
    url = f"{OLLAMA_BASE_URL}/api/embeddings"
    data = {"model": EMBEDDING_MODEL, "prompt": text[:8000]}  # 截断超长文本
    try:
        res = requests.post(url, json=data)
        return res.json()["embedding"]
    except:
        print("向量生成失败，请检查 Ollama 是否启动")
        return []


# ====================== 3. 加载本地知识库 ======================
def load_knowledge_base(folder_path: str = "./media/uploads") -> list:
    """
    加载指定文件夹下所有 txt/md/pdf 作为知识库
    返回：[ {"content": "内容", "embedding": [向量]} ... ]
    """
    knowledge = []
    if not os.path.exists(folder_path):
        os.makedirs(folder_path, exist_ok=True)
        print(f"已创建知识库文件夹：{folder_path}，请放入文件后重试")
        return knowledge

    # 遍历文件夹所有文件
    for filename in os.listdir(folder_path):
        path = os.path.join(folder_path, filename)
        content = get_file_content(path)
        if len(content) > 10:  # 过滤空文件
            embed = get_embedding(content)
            knowledge.append({
                "filename": filename,
                "content": content,
                "embedding": embed
            })
            print(f"已加载：{filename}")
    print(f"\n知识库加载完成，共 {len(knowledge)} 个文件\n")
    return knowledge


# ====================== 4. 从知识库检索相关内容 ======================
def search_knowledge(question: str, knowledge: list, top_k: int = 2) -> str:
    """
    用问题向量匹配知识库，返回最相关的内容
    """
    if not knowledge:
        return ""

    # 1. 生成问题的向量
    q_embed = get_embedding(question)
    if not q_embed:
        return ""

    # 2. 计算余弦相似度
    scores = []
    for item in knowledge:
        k_embed = item["embedding"]
        if not k_embed:
            continue
        sim = cosine_similarity([q_embed], [k_embed])[0][0]
        scores.append((sim, item["content"]))

    # 3. 按相似度排序，取 top_k 条
    scores.sort(reverse=True, key=lambda x: x[0])
    related_text = "\n---\n".join([s[1] for s in scores[:top_k]])
    return related_text


# ====================== 5. 向 Ollama 提问（基于知识库） ======================
def ask_ollama(question: str, knowledge: list, chart_type: str = '') -> str:
    """
    核心：检索知识库 → 拼接提示词 → 发送给 Ollama 回答
    """
    # 1. 检索相关知识
    context = search_knowledge(question, knowledge)
    if not context:
        context = "没有找到相关知识库内容"

    # 2. 根据图表类型生成对应的提示指令
    chart_instructions = {
        'line': '请使用ECharts生成一个**折线图**（line chart），用合理的示例数据展示趋势变化，图表要完整可运行。',
        'bar': '请使用ECharts生成一个**柱状图**（bar chart），用合理的示例数据进行对比展示，图表要完整可运行。',
        'pie': '请使用ECharts生成一个**饼图**（pie chart），用合理的示例数据展示占比分布，图表要完整可运行。',
        'scatter': '请使用ECharts生成一个**散点图**（scatter chart），用合理的示例数据展示数据点分布关系，图表要完整可运行。',
        'radar': '请使用ECharts生成一个**数据雷达图**（radar chart），用合理的示例数据展示多维度指标对比，图表要完整可运行。',
        'surface3d': '请使用ECharts GL生成一个**3D函数图像**（3D surface chart），展示一个三维曲面/函数图像，包含合理的数学函数数据，图表要完整可运行。',
        'bar3d': '请使用ECharts GL生成一个**3D柱状图**（3D bar chart），用合理的示例数据在三维空间中展示柱状对比，图表要完整可运行。',
        'heatmap': '请使用ECharts生成一个**矩阵热力图**（heatmap），用合理的示例数据展示矩阵中数据的强度分布，图表要完整可运行。',
        'comparison-table': '请使用HTML+CSS生成一个**二维对比表格**（三线表样式，仅使用三条横线：顶线、栏目线、底线），用于对比两个知识点的概念和优缺点，表格要完整可直接使用。',
        'tree': '请使用ECharts生成一个**树图**（tree diagram，类似思维导图），用合理的示例数据展示层级结构关系，图表要完整可运行。',
        'graph': '请使用ECharts生成一个**关系图**（graph/知识图谱），用合理的示例数据展示节点之间的关系网络，图表要完整可运行。',
        'flowchart': '请使用Mermaid语法生成一个**流程图**（flowchart），用合理的示例数据展示流程步骤和分支逻辑，图表要完整可运行。在返回的HTML中使用`<div class="mermaid">...</div>`包裹Mermaid代码，并引入Mermaid CDN脚本。',
    }

    chart_prompt = chart_instructions.get(chart_type, '请使用ECharts展示相关图表，最好用3D图表，图表要完整可运行，不要有任何空图。')

    # 3. 构造 RAG 提示词（关键！让模型只根据知识库回答）
    prompt = f"""
你是一个基于知识库的智能助手，请**只根据下面的知识库内容**回答问题。
如果知识库中没有答案，请直接回答："抱歉，我的知识库中没有相关信息"。
返回使用html格式，并融合知识库内容。

{chart_prompt}

=== 知识库内容 ===
{context}
==================

用户问题：{question}
请回答：
    """

    # 4. 调用 Ollama 聊天接口
    url = f"{OLLAMA_BASE_URL}/api/generate"
    data = {
        "model": CHAT_MODEL,
        "prompt": prompt,
        "stream": False,
        "temperature": 0.1  # 低温度=更准确
    }

    try:
        res = requests.post(url, json=data)
        return res.json()["response"].strip()
    except Exception as e:
        return f"请求失败：{str(e)}"




def load_ask(ask, chart_type='', folder_path='./media/uploads'):
    kb = load_knowledge_base(folder_path)
    answer = ask_ollama(ask, kb, chart_type)
    return answer


# ====================== 6. 关键词提取（基于 LLM 主题提取 + 文件名回退） ======================

def _keywords_from_text(text: str) -> list:
    """n-gram 词频统计回退方案"""
    if not text:
        return []
    chinese_pattern = re.compile(r'[一-鿿]{2,4}')
    words = chinese_pattern.findall(text)
    stop_words = {
        '可以', '使用', '一个', '这个', '那个', '我们', '他们', '它们',
        '进行', '没有', '并且', '因为', '所以', '但是', '如果', '虽然',
        '不是', '就是', '还是', '只是', '而且', '或者', '然后', '之后',
        '什么', '怎么', '怎样', '为什么', '哪里', '哪个', '如何',
        '一些', '一种', '一下', '一定', '一样', '一切',
        '这些', '那些', '这里', '那里', '其中', '其他',
        '关于', '对于', '根据', '按照', '通过', '经过',
        '已经', '正在', '将要', '可能', '应该', '需要',
        '的', '了', '在', '是', '有', '和', '与', '或',
        '不', '也', '都', '就', '还', '要', '会', '能',
        '对', '从', '到', '上', '下', '中', '去', '来',
        '这', '那', '其', '之', '以', '及', '而', '等',
        '为', '所', '被', '把', '让', '给', '向', '将',
        '大', '小', '多', '少', '新', '旧', '高', '低',
    }
    freq = {}
    for w in words:
        if w not in stop_words:
            freq[w] = freq.get(w, 0) + 1
    sorted_items = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [{"name": k, "value": v} for k, v in sorted_items[:10]]


def _keywords_from_filename(filename: str) -> list:
    """从文件名中提取关键词（LLM 失败时的回退方案）"""
    name = os.path.splitext(filename)[0]
    # 按常见分隔符切分
    fragments = re.split(r'[-_—·\s【】\[\]（）(),，、。\.]+', name)
    keywords = []
    for f in fragments:
        f = f.strip()
        # 至少2个中文字符，排除纯数字/纯英文
        if len(f) >= 2 and re.search(r'[一-鿿]', f):
            keywords.append({"name": f, "value": 1})
    return keywords


def _llm_filter_keywords(keywords: list, content: str) -> list:
    """调用 Ollama 从候选词中筛选主题关键词，返回保留的 name 列表"""
    if not keywords or len(keywords) <= 3:
        return [kw["name"] for kw in keywords]

    candidate_list = [kw["name"] for kw in keywords]
    prompt = f"""你是教育内容主题分析专家。以下是从一篇文章中提取的候选词组及其出现频率。
请筛选出其中真正代表文章**主题/知识点/核心概念**的关键词（2-6字），排除：
- 无意义的数字/时间/量词组合
- 通用描述性词组（如"非常重"、"比较大"）
- 不是知识概念的日常用语

候选词组：{json.dumps(candidate_list, ensure_ascii=False)}
文章片段：{content[:800]}

严格返回 JSON 数组，只包含应保留的关键词名称，不要其他内容。
示例：["关键词1", "关键词2", "关键词3"]"""

    try:
        url = f"{OLLAMA_BASE_URL}/api/generate"
        data = {
            "model": CHAT_MODEL,
            "prompt": prompt,
            "stream": False,
            "temperature": 0.1
        }
        res = requests.post(url, json=data, timeout=20)
        raw = res.json()["response"].strip()
        raw = re.sub(r'```(?:json)?\s*', '', raw)
        raw = re.sub(r'```', '', raw).strip()
        filtered = json.loads(raw)
        if isinstance(filtered, list) and len(filtered) > 0:
            return [str(item) for item in filtered]
    except Exception as e:
        print(f"LLM 关键词筛选失败: {e}")

    return [kw["name"] for kw in keywords]


def extract_topic_keywords(content: str, filename: str = "") -> list:
    """
    n-gram 词频统计 + LLM 主题筛选，每篇最多 5 个关键词
    返回：[{"name": "关键词", "value": 词频}, ...]（value 是该词在本文中的 n-gram 频次）
    """
    if not content or len(content) < 10:
        result = _keywords_from_filename(filename)
        return [{"name": kw["name"], "value": 1} for kw in result[:5]]

    # 1. n-gram 词频统计，取前 30 个候选
    candidates = _keywords_from_text(content)
    if not candidates:
        result = _keywords_from_filename(filename)
        return [{"name": kw["name"], "value": 1} for kw in result[:5]]

    # 构建 name → frequency 映射（保留原始 n-gram 频次）
    freq_map = {kw["name"]: kw["value"] for kw in candidates}
    candidates = candidates[:30]

    # 2. 用大模型筛选主题关键词
    kept_names = set(_llm_filter_keywords(candidates, content))

    # 3. 保留被筛选出的关键词（使用原始 n-gram 频次），最多 5 个
    result = [{"name": kw["name"], "value": freq_map.get(kw["name"], 1)} for kw in candidates if kw["name"] in kept_names][:5]

    # 4. 文件名关键词补充（仍限制总数 ≤5）
    if len(result) < 5:
        filename_kw = _keywords_from_filename(filename)
        existing_names = {kw["name"] for kw in result}
        for kw in filename_kw:
            if kw["name"] not in existing_names and len(result) < 5:
                result.append({"name": kw["name"], "value": 1})

    return result


def analyze_all_files(folder_path: str = "./media/uploads") -> dict:
    """
    统计所有文件的关键词频次（n-gram 词频跨文件累加）
    每篇最多提取 5 个主题关键词，value = 该词在各文件中的 n-gram 频次总和
    返回：{"top6": [...], "others": [...]}
    """
    if not os.path.exists(folder_path):
        os.makedirs(folder_path, exist_ok=True)
        return {"top6": [], "others": []}

    merged_freq = {}

    for filename in os.listdir(folder_path):
        path = os.path.join(folder_path, filename)
        if not os.path.isfile(path):
            continue
        content = get_file_content(path)
        if len(content) > 10:
            keywords = extract_topic_keywords(content, filename)
            for kw in keywords:
                name = kw["name"]
                merged_freq[name] = merged_freq.get(name, 0) + kw["value"]

    sorted_all = sorted(merged_freq.items(), key=lambda x: x[1], reverse=True)
    all_keywords = [{"name": k, "value": v} for k, v in sorted_all]

    top6 = all_keywords[:6]
    others = all_keywords[6:]

    return {"top6": top6, "others": others}


# ====================== 主程序：启动问答 ======================
if __name__ == "__main__":
    print("启动本地知识库问答系统...")

    # 1. 加载知识库（自动读取 ./knowledge 文件夹）
    kb = load_knowledge_base()

    # 2. 循环问答
    while True:
        user_q = input("\n请输入你的问题（输入 exit 退出）：")
        if user_q.lower() == "exit":
            print("退出程序")
            break

        print("正在检索知识库并生成答案...")
        answer = ask_ollama(user_q, kb)
        print(f"\n助手回答：\n{answer}")